# extractor.py
#
# Unified extraction layer.
# Accepts: PDF, DOCX, image (.png/.jpg/.tiff), CSV/TXT/DAT
# Returns a single `data` dict consumed by all downstream components.
#
# v2 additions:
#   - Image detection: count + basic analysis (grayscale, color diversity)
#   - Table detection: headers, element columns, percentage patterns

import os
import re
import logging
import numpy as np

logger = logging.getLogger("vidur.extractor")

# ── periodic table symbols used for element-column detection ─────────────────
_ALL_ELEMENTS = {
    "h","he","li","be","b","c","n","o","f","ne","na","mg","al","si","p","s",
    "cl","ar","k","ca","sc","ti","v","cr","mn","fe","co","ni","cu","zn","ga",
    "ge","as","se","br","kr","rb","sr","y","zr","nb","mo","tc","ru","rh","pd",
    "ag","cd","in","sn","sb","te","i","xe","cs","ba","la","ce","pr","nd","pm",
    "sm","eu","gd","tb","dy","ho","er","tm","yb","lu","hf","ta","w","re","os",
    "ir","pt","au","hg","tl","pb","bi","po","at","rn","fr","ra","ac","th","pa",
    "u","np","pu",
}

_PERCENTAGE_HEADERS = {"wt%", "at%", "weight %", "atomic %", "wt. %", "at. %",
                        "weight%", "atomic%", "mass%", "mol%"}


def extract(file_path: str) -> dict:
    """
    Extract text + numeric signals + image signals + table signals.

    Returns:
        {
            # ── core (unchanged) ──────────────────────────────────────────
            "file_path":    str,
            "filename":     str,
            "extension":    str,
            "text":         str,
            "numeric_data": np.ndarray | None,
            "magic_bytes":  bytes,
            "metadata":     dict,

            # ── Upgrade A: image signals ──────────────────────────────────
            "has_images":        bool,
            "image_count":       int,
            "image_dimensions":  list[tuple],   # [(w, h), ...]
            "image_signals": {
                "is_grayscale":        bool,   # standalone image or majority grayscale
                "low_color_diversity": bool,   # unique colours / total pixels < threshold
                "microscopy_like":     bool,   # combined heuristic flag
            },

            # ── Upgrade B: table signals ──────────────────────────────────
            "has_table":     bool,
            "table_headers": list[str],
            "table_signals": {
                "has_element_columns":    bool,
                "has_percentage_headers": bool,
                "percentages_sum_100":    bool,
                "element_columns":        list[str],
            },
        }
    """
    ext      = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)

    result = {
        # core
        "file_path":    file_path,
        "filename":     filename,
        "extension":    ext,
        "text":         "",
        "numeric_data": None,
        "magic_bytes":  b"",
        "metadata":     {"source_format": ext.lstrip(".")},
        # image signals
        "has_images":       False,
        "image_count":      0,
        "image_dimensions": [],
        "image_signals": {
            "is_grayscale":        False,
            "low_color_diversity": False,
            "microscopy_like":     False,
        },
        # table signals
        "has_table":     False,
        "table_headers": [],
        "table_signals": {
            "has_element_columns":    False,
            "has_percentage_headers": False,
            "percentages_sum_100":    False,
            "element_columns":        [],
        },
    }

    # Always read magic bytes first
    try:
        with open(file_path, "rb") as f:
            result["magic_bytes"] = f.read(512)
    except Exception as e:
        logger.warning(f"Could not read magic bytes from {filename}: {e}")
        return result

    # Dispatch by extension
    if ext == ".pdf":
        _extract_pdf(file_path, result)
    elif ext in (".docx", ".doc"):
        _extract_docx(file_path, result)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"):
        _extract_image(file_path, result)
    elif ext in (".csv", ".txt", ".dat", ".asc", ".xy", ".dsp", ".abs",
                 ".emsa", ".msa"):
        _extract_ascii(file_path, result)
    else:
        _extract_ascii(file_path, result, strict=False)

    logger.debug(
        f"Extracted [{filename}]: "
        f"text_len={len(result['text'])}, "
        f"images={result['image_count']}, "
        f"has_table={result['has_table']}, "
        f"numeric_shape={result['numeric_data'].shape if result['numeric_data'] is not None else None}"
    )
    return result


# ── format-specific extractors ────────────────────────────────────────────────

def _extract_pdf(path: str, result: dict):
    """Extract text + count embedded images from PDF."""
    try:
        import pypdf
        reader = pypdf.PdfReader(path)
        pages  = []
        images = []

        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pass
            # Count embedded images
            try:
                for img_obj in page.images:
                    images.append(img_obj)
            except Exception:
                pass

        text = " ".join(pages)
        result["text"]                    = text.lower()
        result["metadata"]["page_count"]  = len(reader.pages)
        result["has_images"]              = len(images) > 0
        result["image_count"]             = len(images)

        # Extract dimensions where available
        for img in images:
            try:
                result["image_dimensions"].append((img.width, img.height))
            except Exception:
                pass

        _analyse_embedded_images(images, result)

    except ImportError:
        logger.warning("pypdf not installed — PDF extraction unavailable.")
    except Exception as e:
        logger.warning(f"PDF extraction failed for {result['filename']}: {e}")


def _extract_docx(path: str, result: dict):
    """Extract text, structured tables, and count embedded images from DOCX."""
    try:
        import docx
        doc   = docx.Document(path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]

        # ── structured table extraction ───────────────────────────────────
        all_table_rows = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_rows.append(cells)
                    lines.append(" ".join(cells))
            if table_rows:
                all_table_rows.extend(table_rows)

        if all_table_rows:
            _analyse_table_rows(all_table_rows, result)

        # ── image counting ────────────────────────────────────────────────
        # python-docx exposes inline shapes
        try:
            img_count = len(doc.inline_shapes)
            result["has_images"]  = img_count > 0
            result["image_count"] = img_count
        except Exception:
            pass

        result["text"]                        = " ".join(lines).lower()
        result["metadata"]["paragraph_count"] = len(doc.paragraphs)

    except ImportError:
        logger.warning("python-docx not installed — DOCX extraction unavailable.")
    except Exception as e:
        logger.warning(f"DOCX extraction failed for {result['filename']}: {e}")


def _extract_image(path: str, result: dict):
    """Analyse a standalone image file for SEM-like properties."""
    # Filename as text signal
    result["text"] = result["filename"].lower().replace("_", " ").replace("-", " ")
    result["has_images"]  = True
    result["image_count"] = 1

    try:
        from PIL import Image
        with Image.open(path) as img:
            w, h = img.size
            result["metadata"]["image_size"]   = (w, h)
            result["metadata"]["image_mode"]   = img.mode
            result["metadata"]["image_format"] = img.format
            result["image_dimensions"]         = [(w, h)]

            sigs = result["image_signals"]

            # Grayscale check
            if img.mode in ("L", "LA", "P"):
                sigs["is_grayscale"] = True
            elif img.mode in ("RGB", "RGBA"):
                # Convert and compare channels
                arr = np.array(img.convert("RGB"))
                r, g, b = arr[:, :, 0], arr[:, :, 1], arr[:, :, 2]
                channel_diff = (np.abs(r.astype(int) - g.astype(int)).mean() +
                                np.abs(g.astype(int) - b.astype(int)).mean())
                if channel_diff < 8.0:   # channels nearly identical → grayscale-like
                    sigs["is_grayscale"] = True

            # Color diversity: sample at most 10k pixels
            arr_rgb = np.array(img.convert("RGB"))
            flat    = arr_rgb.reshape(-1, 3)
            if len(flat) > 10_000:
                idx  = np.random.choice(len(flat), 10_000, replace=False)
                flat = flat[idx]
            unique_colors = len(np.unique(flat, axis=0))
            diversity     = unique_colors / len(flat)
            if diversity < 0.05:          # < 5 % unique colours → low diversity
                sigs["low_color_diversity"] = True

            # Combined SEM heuristic
            sigs["microscopy_like"] = sigs["is_grayscale"] or sigs["low_color_diversity"]

    except ImportError:
        logger.debug("Pillow not installed — image analysis skipped.")
    except Exception as e:
        logger.warning(f"Image analysis failed for {result['filename']}: {e}")


def _extract_ascii(path: str, result: dict, strict: bool = True):
    """
    Extract text lines + numeric table + structured table signals.
    """
    try:
        with open(path, "r", errors="replace") as f:
            raw_lines = f.readlines()
    except Exception as e:
        if strict:
            logger.warning(f"Cannot read {result['filename']} as ASCII: {e}")
        return

    text_lines = [l.strip() for l in raw_lines if l.strip()]
    result["text"] = " ".join(text_lines).lower()

    # Numeric table
    numeric = _parse_numeric_table(raw_lines)
    if numeric is not None and numeric.shape[0] > 1:
        result["numeric_data"] = numeric

    # Structured table signals from raw lines
    _analyse_ascii_table(raw_lines, result)


# ── image analysis helpers ────────────────────────────────────────────────────

def _analyse_embedded_images(images: list, result: dict):
    """
    Analyse embedded PDF images for SEM-like properties using Pillow.
    Works on pypdf ImageFile objects which expose .data bytes.
    """
    sigs = result["image_signals"]
    if not images:
        return

    grayscale_count = 0
    low_diversity_count = 0

    for img_obj in images[:5]:   # inspect first 5 images max
        try:
            import io
            from PIL import Image
            pil_img = Image.open(io.BytesIO(img_obj.data))

            if pil_img.mode in ("L", "LA", "P"):
                grayscale_count += 1
            elif pil_img.mode in ("RGB", "RGBA"):
                arr = np.array(pil_img.convert("RGB"))
                r, g, b = arr[:, :, 0], arr[:, :, 1], arr[:, :, 2]
                channel_diff = (np.abs(r.astype(int) - g.astype(int)).mean() +
                                np.abs(g.astype(int) - b.astype(int)).mean())
                if channel_diff < 8.0:
                    grayscale_count += 1

                flat = arr.reshape(-1, 3)
                if len(flat) > 10_000:
                    flat = flat[np.random.choice(len(flat), 10_000, replace=False)]
                diversity = len(np.unique(flat, axis=0)) / len(flat)
                if diversity < 0.05:
                    low_diversity_count += 1
        except Exception:
            continue

    inspected = min(len(images), 5)
    if inspected > 0:
        if grayscale_count / inspected >= 0.5:
            sigs["is_grayscale"] = True
        if low_diversity_count / inspected >= 0.5:
            sigs["low_color_diversity"] = True
        sigs["microscopy_like"] = sigs["is_grayscale"] or sigs["low_color_diversity"]


# ── table analysis helpers ────────────────────────────────────────────────────

def _analyse_table_rows(rows: list[list[str]], result: dict):
    """
    Analyse a list of table rows (each row = list of cell strings).
    Looks for element columns and percentage headers.
    """
    if not rows:
        return

    result["has_table"] = True
    tsigs = result["table_signals"]

    # Treat first row as potential header
    headers = [c.strip().lower() for c in rows[0]]
    result["table_headers"] = headers

    # Check for percentage headers
    for h in headers:
        if h in _PERCENTAGE_HEADERS or any(ph in h for ph in _PERCENTAGE_HEADERS):
            tsigs["has_percentage_headers"] = True
            break

    # Check for element symbol columns (headers that are pure element symbols)
    elem_cols = [h for h in headers if h in _ALL_ELEMENTS]
    if elem_cols:
        tsigs["has_element_columns"] = True
        tsigs["element_columns"]     = elem_cols

    # Check if any data column sums to ~100 (composition percentage)
    if len(rows) > 1:
        _check_percentages_sum(rows[1:], headers, tsigs)


def _analyse_ascii_table(raw_lines: list[str], result: dict):
    """
    Detect and analyse tabular structure in ASCII/CSV files.
    Finds a header row and checks for element/percentage columns.
    """
    # Detect the most likely delimiter
    delimiter = _detect_delimiter(raw_lines)

    header_row  = None
    data_rows   = []
    found_header = False

    for line in raw_lines:
        line = line.strip()
        if not line or line[0] in "#!;$":
            continue
        parts = [p.strip() for p in line.split(delimiter) if p.strip()]
        if len(parts) < 2:
            continue

        # Is this a header row? Heuristic: has non-numeric tokens
        if not found_header:
            non_numeric = sum(1 for p in parts if not _is_number(p))
            if non_numeric >= 1:
                header_row   = parts
                found_header = True
                continue

        # Collect numeric data rows
        if found_header:
            if all(_is_number(p) for p in parts):
                data_rows.append(parts)

    if header_row is None:
        return

    result["has_table"]     = True
    result["table_headers"] = [h.lower() for h in header_row]
    tsigs = result["table_signals"]

    headers_lower = [h.lower() for h in header_row]

    # Percentage headers
    for h in headers_lower:
        if h in _PERCENTAGE_HEADERS or any(ph in h for ph in _PERCENTAGE_HEADERS):
            tsigs["has_percentage_headers"] = True
            break

    # Element symbol columns
    elem_cols = [h for h in headers_lower if h in _ALL_ELEMENTS]
    if elem_cols:
        tsigs["has_element_columns"] = True
        tsigs["element_columns"]     = elem_cols

    # Check if data columns sum to ~100
    if data_rows:
        _check_percentages_sum_ascii(data_rows, headers_lower, tsigs)


def _check_percentages_sum(data_rows: list[list[str]], headers: list[str],
                            tsigs: dict):
    """Check if any numeric column in structured table rows sums to ~100."""
    for col_idx in range(len(headers)):
        vals = []
        for row in data_rows:
            if col_idx < len(row):
                try:
                    vals.append(float(row[col_idx]))
                except (ValueError, TypeError):
                    pass
        if vals:
            col_sum = sum(vals)
            if 90 <= col_sum <= 110:
                tsigs["percentages_sum_100"] = True
                return


def _check_percentages_sum_ascii(data_rows: list[list[str]], headers: list[str],
                                  tsigs: dict):
    """Check if any numeric column in ASCII rows sums to ~100."""
    if not data_rows:
        return
    n_cols = max(len(row) for row in data_rows)
    for col_idx in range(n_cols):
        vals = []
        for row in data_rows:
            if col_idx < len(row):
                try:
                    vals.append(float(row[col_idx]))
                except ValueError:
                    pass
        if vals:
            col_sum = sum(vals)
            if 90 <= col_sum <= 110:
                tsigs["percentages_sum_100"] = True
                return


# ── shared utilities ──────────────────────────────────────────────────────────

def _detect_delimiter(lines: list[str]) -> str:
    """Heuristically detect the most common delimiter in a list of lines."""
    counts = {",": 0, "\t": 0, " ": 0}
    for line in lines[:20]:
        counts[","]  += line.count(",")
        counts["\t"] += line.count("\t")
    if counts[","] >= counts["\t"] and counts[","] > 0:
        return ","
    if counts["\t"] > 0:
        return "\t"
    return None   # space/auto


def _is_number(s: str) -> bool:
    try:
        float(s)
        return True
    except (ValueError, TypeError):
        return False


def _parse_numeric_table(lines: list) -> np.ndarray | None:
    """
    Try to parse a 2D numeric table from raw text lines.
    Handles CSV (comma), TSV (tab), or space-delimited. Auto-skips headers.
    """
    import io
    for delimiter in (None, ",", "\t"):
        for skip in range(min(50, len(lines))):
            try:
                subset = "".join(lines[skip:])
                data   = np.loadtxt(
                    io.StringIO(subset),
                    delimiter=delimiter,
                    comments=["#", "!", ";", "$"]
                )
                if data.ndim == 1:
                    data = data.reshape(-1, 1)
                if data.ndim == 2 and data.shape[0] > 1 and data.shape[1] >= 1:
                    return data
            except Exception:
                continue
    return None
